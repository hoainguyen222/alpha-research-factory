"""
====================================================================================================
MODULE: Office Documents & Structured Data Extractor
FILE: extractors/office_extractor.py
====================================================================================================
CHỨC NĂNG (YÊU CẦU 1 & 2):
1. Trích xuất toàn diện các file văn phòng và bảng tính:
   - Word (.docx, .doc): Toàn bộ đoạn văn, tiêu đề phân cấp, bảng biểu dưới dạng Markdown Table.
   - Excel (.xlsx, .xls): Đọc tất cả các Sheets, trích xuất cấu trúc cột, thống kê mô tả, bảng Markdown.
   - CSV / TSV (.csv, .tsv): Tự động nhận diện dấu phân cách, trích xuất Schema và dữ liệu mẫu.
   - JSON / Text / Markdown (.json, .txt, .md): Đọc và định dạng cấu trúc rõ ràng.
====================================================================================================
"""

import io
import json
from pathlib import Path
from typing import Dict, Any, Optional, Union, List
import pandas as pd
import docx

from bypass.file_decryptor import FileDecryptor


class OfficeExtractor:
    """Module trích xuất các file Word, Excel, CSV, JSON, Text."""

    @classmethod
    def extract_docx(cls, file_input: Union[str, Path, bytes], filename: str = "document.docx") -> Dict[str, Any]:
        """Trích xuất file Word .docx."""
        try:
            if isinstance(file_input, bytes):
                stream = io.BytesIO(file_input)
                doc = docx.Document(stream)
            else:
                doc = docx.Document(str(file_input))

            lines = []
            title = Path(filename).stem
            first_heading_found = False

            # 1. Trích xuất các đoạn văn
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                # Định dạng tiêu đề theo style
                style_name = p.style.name.lower() if p.style else ""
                if "heading 1" in style_name or "title" in style_name:
                    lines.append(f"\n# {text}\n")
                    if not first_heading_found:
                        title = text
                        first_heading_found = True
                elif "heading 2" in style_name:
                    lines.append(f"\n## {text}\n")
                elif "heading 3" in style_name:
                    lines.append(f"\n### {text}\n")
                else:
                    lines.append(text)

            # 2. Trích xuất các bảng biểu trong file Word
            if doc.tables:
                lines.append("\n### [TABLES IN DOCUMENT]\n")
                for t_idx, table in enumerate(doc.tables):
                    lines.append(f"**Table {t_idx + 1}:**")
                    table_rows = []
                    for row in table.rows:
                        row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        table_rows.append(row_cells)
                    if table_rows:
                        # Chuyển đổi sang Markdown table
                        header = table_rows[0]
                        lines.append("| " + " | ".join(header) + " |")
                        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                        for r in table_rows[1:]:
                            # Cân bằng số cột
                            padded_r = r + [""] * (len(header) - len(r))
                            lines.append("| " + " | ".join(padded_r[:len(header)]) + " |")
                        lines.append("")

            full_text = "\n\n".join(lines)
            return {
                "title": title,
                "text": full_text,
                "metadata": {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
                "extraction_method": "python-docx Structured Parser",
                "bypass_status": "Direct"
            }

        except Exception as e:
            # Fallback sang phục hồi XML nhị phân
            if isinstance(file_input, bytes):
                salvaged = FileDecryptor.salvage_text_from_corrupted_binary(file_input)
            else:
                with open(file_input, "rb") as f:
                    salvaged = FileDecryptor.salvage_text_from_corrupted_binary(f.read())
            return {
                "title": Path(filename).stem,
                "text": salvaged,
                "metadata": {"salvaged": True, "error": str(e)},
                "extraction_method": "Office XML Salvage Fallback",
                "bypass_status": f"Corrupted / Recovered ({str(e)})"
            }

    @classmethod
    def extract_excel(cls, file_input: Union[str, Path, bytes], filename: str = "spreadsheet.xlsx") -> Dict[str, Any]:
        """Trích xuất file Excel đa sheets (.xlsx, .xls)."""
        try:
            target = io.BytesIO(file_input) if isinstance(file_input, bytes) else str(file_input)
            xls = pd.ExcelFile(target)
            sheet_names = xls.sheet_names

            sections = []
            sections.append(f"# EXCEL WORKBOOK: {Path(filename).name}")
            sections.append(f"**Total Sheets ({len(sheet_names)}):** {', '.join(sheet_names)}\n")

            total_rows_all = 0
            for sname in sheet_names:
                df = pd.read_excel(xls, sheet_name=sname)
                total_rows_all += len(df)
                sections.append(f"## Sheet: {sname}")
                sections.append(f"- **Dimensions:** {df.shape[0]} rows x {df.shape[1]} columns")
                sections.append(f"- **Columns:** {', '.join(str(c) for c in df.columns)}")
                
                # Bảng thống kê mô tả nếu có cột số
                desc = df.describe()
                if not desc.empty:
                    sections.append("\n**Statistical Summary:**")
                    sections.append(desc.to_markdown())

                # Dữ liệu xem trước (top 20 dòng đầu)
                sections.append(f"\n**Data Preview (Top {min(20, len(df))} rows):**")
                sections.append(df.head(20).to_markdown(index=False))
                sections.append("\n---\n")

            return {
                "title": Path(filename).stem,
                "text": "\n".join(sections),
                "metadata": {"sheet_names": sheet_names, "total_sheets": len(sheet_names), "total_rows": total_rows_all},
                "extraction_method": "Pandas Excel Multi-Sheet Engine",
                "bypass_status": "Direct"
            }

        except Exception as e:
            return {
                "title": Path(filename).stem,
                "text": f"[EXCEL EXTRACTION ERROR]: {str(e)}",
                "metadata": {"error": str(e)},
                "extraction_method": "Excel Parser Error",
                "bypass_status": "Failed"
            }

    @classmethod
    def extract_csv(cls, file_input: Union[str, Path, bytes], filename: str = "data.csv") -> Dict[str, Any]:
        """Trích xuất file CSV / TSV."""
        try:
            target = io.BytesIO(file_input) if isinstance(file_input, bytes) else str(file_input)
            
            # Tự động nhận diện định dạng phân cách (comma, tab, semicolon)
            try:
                df = pd.read_csv(target, sep=None, engine='python')
            except Exception:
                if isinstance(target, io.BytesIO):
                    target.seek(0)
                df = pd.read_csv(target, sep=",")

            sections = []
            sections.append(f"# CSV DATASET: {Path(filename).name}")
            sections.append(f"- **Total Rows:** {len(df):,}")
            sections.append(f"- **Total Columns:** {len(df.columns)}")
            sections.append(f"- **Column Names:** {', '.join(str(c) for c in df.columns)}\n")

            # Thống kê mô tả
            desc = df.describe()
            if not desc.empty:
                sections.append("## Statistical Summary")
                sections.append(desc.to_markdown())
                sections.append("")

            # Preview bảng Markdown (30 dòng đầu)
            sections.append(f"## Data Preview (First {min(30, len(df))} rows)")
            sections.append(df.head(30).to_markdown(index=False))

            return {
                "title": Path(filename).stem,
                "text": "\n".join(sections),
                "metadata": {"rows": len(df), "columns": len(df.columns), "column_list": list(df.columns)},
                "extraction_method": "Pandas CSV Auto-Delimiter Parser",
                "bypass_status": "Direct"
            }

        except Exception as e:
            return {
                "title": Path(filename).stem,
                "text": f"[CSV EXTRACTION ERROR]: {str(e)}",
                "metadata": {"error": str(e)},
                "extraction_method": "CSV Parser Error",
                "bypass_status": "Failed"
            }

    @classmethod
    def extract_text_or_json(cls, file_input: Union[str, Path, bytes], filename: str = "file.txt") -> Dict[str, Any]:
        """Trích xuất file TXT, Markdown, JSON."""
        try:
            if isinstance(file_input, bytes):
                raw_str = file_input.decode("utf-8", errors="replace")
            else:
                with open(file_input, "r", encoding="utf-8", errors="replace") as f:
                    raw_str = f.read()

            title = Path(filename).stem
            # Nếu là JSON, format lại cho đẹp
            if filename.lower().endswith(".json"):
                try:
                    obj = json.loads(raw_str)
                    formatted_text = f"# JSON DATA: {filename}\n```json\n" + json.dumps(obj, indent=2, ensure_ascii=False) + "\n```"
                    return {
                        "title": title,
                        "text": formatted_text,
                        "metadata": {"is_json": True, "type": type(obj).__name__},
                        "extraction_method": "JSON Structured Formatter",
                        "bypass_status": "Direct"
                    }
                except Exception:
                    pass

            return {
                "title": title,
                "text": raw_str,
                "metadata": {"length_chars": len(raw_str)},
                "extraction_method": "Raw Text Reader",
                "bypass_status": "Direct"
            }

        except Exception as e:
            return {
                "title": Path(filename).stem,
                "text": f"[TEXT READ ERROR]: {str(e)}",
                "metadata": {"error": str(e)},
                "extraction_method": "Text Reader Error",
                "bypass_status": "Failed"
            }
