"""
====================================================================================================
MODULE: Multi-Format Entry Exporter (Word, Excel, PDF, CSV, JSON, Text)
FILE: vault/entry_exporter.py
====================================================================================================
CHỨC NĂNG:
1. Xuất dữ liệu của từng bản ghi nghiên cứu thành các định dạng phổ biến:
   - Word (.docx): Tài liệu định dạng chuẩn với Tiêu đề, Metadata, Tóm tắt (Summary), Nội dung CTX đầy đủ, Bảng biểu.
   - Excel (.xlsx): Sổ làm việc Excel với 2 sheet (Overview & Metadata, Full_Content_Lines).
   - PDF (.pdf): Tài liệu PDF chuyên nghiệp tạo bằng ReportLab kèm trang bìa/tóm tắt và nội dung.
   - CSV (.csv): File CSV chứa đầy đủ các trường (ID, Title, Type, Summary, CTX, URL, Date).
   - JSON (.json): File JSON có cấu trúc đẹp mắt (indent=2) hỗ trợ máy đọc và API pipeline.
   - Text (.txt): Văn bản thô chuẩn hóa.
2. Trả về bytes dữ liệu kèm mime-type và filename an toàn để tải xuống trực tiếp.
====================================================================================================
"""

import io
import json
import re
from pathlib import Path
from typing import Dict, Any, Tuple
import pandas as pd
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
from reportlab.pdfgen import canvas


class EntryExporter:
    """Module chuyển đổi và xuất bản ghi nghiên cứu sang mọi định dạng cơ bản."""

    @staticmethod
    def _sanitize_filename(title: str, max_len: int = 40) -> str:
        clean = re.sub(r'[\\/*?:"<>|]', "", title or "research")
        clean = re.sub(r'[\s\t\n\r]+', "_", clean).strip("_")
        return (clean[:max_len] or "research_entry").rstrip("_")

    @classmethod
    def export_to_docx(cls, entry: Dict[str, Any]) -> Tuple[bytes, str, str]:
        """Xuất bản ghi sang file Word (.docx)."""
        doc = docx.Document()

        # Thiết lập lề trang
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        entry_id = entry.get("id") or entry.get("ID", "RES-0000")
        title = entry.get("title") or entry.get("TITLE", "Untitled Research")
        source_type = entry.get("type") or entry.get("TYPE", "OTHER")
        note = entry.get("note") or entry.get("NOTE", "")
        ctx = entry.get("ctx") or entry.get("CTX", "")
        web = entry.get("web") or entry.get("WEB", "")
        created_at = entry.get("created_at") or entry.get("CREATED_AT", "")

        # 1. Tiêu đề chính
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(f"[{source_type}] {title}")
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(30, 41, 59)
        title_p.paragraph_format.space_after = Pt(12)

        # 2. Bảng Metadata
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = 'Light Shading Accent 1' if 'Light Shading Accent 1' in [s.name for s in doc.styles] else 'Table Grid'
        
        meta_rows = [
            ("Mã Định Danh (ID):", str(entry_id)),
            ("Phân Loại Nguồn (Type):", str(source_type)),
            ("Nguồn / Đường Dẫn (URL):", str(web)),
            ("Thời Gian Lưu Trữ:", str(created_at).replace("T", " ").split(".")[0])
        ]
        for idx, (label, val) in enumerate(meta_rows):
            meta_table.rows[idx].cells[0].paragraphs[0].text = label
            meta_table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
            meta_table.rows[idx].cells[1].paragraphs[0].text = val

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        # 3. Phần Tóm Tắt & Ghi Chú Chính (Executive Summary)
        if note:
            h_sum = doc.add_heading("1. Tóm Tắt Nội Dung Chính & Ghi Chú Phân Tích (Executive Summary)", level=1)
            h_sum.paragraph_format.space_before = Pt(14)
            h_sum.paragraph_format.space_after = Pt(6)
            
            note_p = doc.add_paragraph()
            note_p.paragraph_format.left_indent = Inches(0.2)
            note_run = note_p.add_run(note)
            note_run.font.size = Pt(11)
            note_run.font.italic = True

        # 4. Phần Nội Dung Đầy Đủ (Full Extracted Context)
        h_ctx = doc.add_heading("2. Toàn Bộ Nội Dung Chi Tiết Trích Xuất (Full Context)", level=1)
        h_ctx.paragraph_format.space_before = Pt(16)
        h_ctx.paragraph_format.space_after = Pt(8)

        for paragraph_text in (ctx or "").split("\n\n"):
            clean_p = paragraph_text.strip()
            if not clean_p:
                continue
            if clean_p.startswith("# "):
                doc.add_heading(clean_p.lstrip("# ").strip(), level=2)
            elif clean_p.startswith("## "):
                doc.add_heading(clean_p.lstrip("# ").strip(), level=3)
            elif clean_p.startswith("### "):
                doc.add_heading(clean_p.lstrip("# ").strip(), level=4)
            else:
                p = doc.add_paragraph(clean_p)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(6)

        buf = io.BytesIO()
        doc.save(buf)
        filename = f"{entry_id}_{cls._sanitize_filename(title)}.docx"
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename

    @classmethod
    def export_to_excel(cls, entry: Dict[str, Any]) -> Tuple[bytes, str, str]:
        """Xuất bản ghi sang file Excel (.xlsx) đa sheet."""
        entry_id = entry.get("id") or entry.get("ID", "RES-0000")
        title = entry.get("title") or entry.get("TITLE", "Untitled")
        source_type = entry.get("type") or entry.get("TYPE", "OTHER")
        note = entry.get("note") or entry.get("NOTE", "")
        ctx = entry.get("ctx") or entry.get("CTX", "")
        web = entry.get("web") or entry.get("WEB", "")
        created_at = entry.get("created_at") or entry.get("CREATED_AT", "")

        # Sheet 1: Metadata & Summary
        df_overview = pd.DataFrame([
            {"Thuộc tính (Field)": "ID Bản Ghi", "Giá trị (Value)": entry_id},
            {"Thuộc tính (Field)": "Tiêu Đề (Title)", "Giá trị (Value)": title},
            {"Thuộc tính (Field)": "Phân Loại (Type)", "Giá trị (Value)": source_type},
            {"Thuộc tính (Field)": "Nguồn Gốc (URL/File)", "Giá trị (Value)": web},
            {"Thuộc tính (Field)": "Thời Gian Tạo", "Giá trị (Value)": created_at},
            {"Thuộc tính (Field)": "Tổng Số Từ", "Giá trị (Value)": len(ctx.split()) if ctx else 0},
            {"Thuộc tính (Field)": "Tóm Tắt & Ghi Chú (Summary)", "Giá trị (Value)": note}
        ])

        # Sheet 2: Content Breakdown (Tách theo từng dòng hoặc đoạn)
        lines = [l.strip() for l in (ctx or "").split("\n") if l.strip()]
        df_content = pd.DataFrame({
            "Line_Number": list(range(1, len(lines) + 1)),
            "Extracted_Text_Content": lines
        })

        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine='openpyxl') as writer:
            df_overview.to_excel(writer, sheet_name="Overview & Summary", index=False)
            df_content.to_excel(writer, sheet_name="Full_Content_Lines", index=False)

        filename = f"{entry_id}_{cls._sanitize_filename(title)}.xlsx"
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", filename

    @classmethod
    def export_to_csv(cls, entry: Dict[str, Any]) -> Tuple[bytes, str, str]:
        """Xuất bản ghi sang file CSV."""
        entry_id = entry.get("id") or entry.get("ID", "RES-0000")
        title = entry.get("title") or entry.get("TITLE", "Untitled")

        df = pd.DataFrame([{
            "ID": entry_id,
            "TITLE": title,
            "TYPE": entry.get("type") or entry.get("TYPE", "OTHER"),
            "SUMMARY_NOTE": entry.get("note") or entry.get("NOTE", ""),
            "FULL_CONTEXT_TEXT": entry.get("ctx") or entry.get("CTX", ""),
            "SOURCE_WEB_URL": entry.get("web") or entry.get("WEB", ""),
            "CREATED_AT": entry.get("created_at") or entry.get("CREATED_AT", ""),
            "RAW_FILE_PATH": entry.get("raw_file_path") or entry.get("RAW_FILE_PATH", "")
        }])

        buf = io.BytesIO()
        df.to_csv(buf, index=False, encoding="utf-8-sig")
        filename = f"{entry_id}_{cls._sanitize_filename(title)}.csv"
        return buf.getvalue(), "text/csv; charset=utf-8", filename

    @classmethod
    def export_to_json(cls, entry: Dict[str, Any]) -> Tuple[bytes, str, str]:
        """Xuất bản ghi sang file JSON có cấu trúc."""
        entry_id = entry.get("id") or entry.get("ID", "RES-0000")
        title = entry.get("title") or entry.get("TITLE", "Untitled")

        # Parse metadata nếu là chuỗi string
        meta = entry.get("metadata") or entry.get("METADATA", {})
        if isinstance(meta, str):
            try:
                meta = json.loads(meta)
            except Exception:
                pass

        clean_record = {
            "ID": entry_id,
            "TITLE": title,
            "TYPE": entry.get("type") or entry.get("TYPE", "OTHER"),
            "SUMMARY_NOTE": entry.get("note") or entry.get("NOTE", ""),
            "CTX": entry.get("ctx") or entry.get("CTX", ""),
            "WEB": entry.get("web") or entry.get("WEB", ""),
            "METADATA": meta,
            "RAW_FILE_PATH": entry.get("raw_file_path") or entry.get("RAW_FILE_PATH", ""),
            "CREATED_AT": entry.get("created_at") or entry.get("CREATED_AT", "")
        }

        json_bytes = json.dumps(clean_record, indent=2, ensure_ascii=False).encode("utf-8")
        filename = f"{entry_id}_{cls._sanitize_filename(title)}.json"
        return json_bytes, "application/json; charset=utf-8", filename

    @classmethod
    def export_to_pdf(cls, entry: Dict[str, Any]) -> Tuple[bytes, str, str]:
        """Xuất bản ghi sang file PDF chuyên nghiệp bằng ReportLab."""
        entry_id = entry.get("id") or entry.get("ID", "RES-0000")
        title = entry.get("title") or entry.get("TITLE", "Untitled Research")
        source_type = entry.get("type") or entry.get("TYPE", "OTHER")
        note = entry.get("note") or entry.get("NOTE", "")
        ctx = entry.get("ctx") or entry.get("CTX", "")
        web = entry.get("web") or entry.get("WEB", "N/A")
        created_at = (entry.get("created_at") or entry.get("CREATED_AT", "")).replace("T", " ").split(".")[0]

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=40,
            rightMargin=40,
            topMargin=40,
            bottomMargin=40
        )

        styles = getSampleStyleSheet()
        
        # Tùy biến styles
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#1e293b"),
            spaceAfter=10
        )
        
        h1_style = ParagraphStyle(
            'SectionH1',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#4338ca"),
            spaceBefore=12,
            spaceAfter=6
        )

        body_style = ParagraphStyle(
            'BodyClean',
            parent=styles['BodyText'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor("#334155"),
            spaceAfter=6
        )

        note_style = ParagraphStyle(
            'NoteBox',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor("#0f172a")
        )

        story = []

        # 1. Header & Title
        story.append(Paragraph(f"<b>[{source_type}] {title}</b>", title_style))
        story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#6366f1"), spaceAfter=10))

        # 2. Metadata Table
        meta_data = [
            [Paragraph("<b>ID Bản Ghi:</b>", body_style), Paragraph(str(entry_id), body_style)],
            [Paragraph("<b>Phân Loại:</b>", body_style), Paragraph(str(source_type), body_style)],
            [Paragraph("<b>Nguồn/URL:</b>", body_style), Paragraph(str(web)[:90], body_style)],
            [Paragraph("<b>Thời Gian:</b>", body_style), Paragraph(str(created_at), body_style)],
        ]
        t = Table(meta_data, colWidths=[110, 400])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

        # 3. Tóm Tắt Nội Dung Chính (Executive Summary)
        if note:
            story.append(Paragraph("1. Tóm Tắt Nội Dung Chính & Ghi Chú (Executive Summary)", h1_style))
            # Hộp nền tóm tắt
            clean_note_html = note.replace("\n", "<br/>").replace("<", "&lt;").replace(">", "&gt;")
            note_table = Table([[Paragraph(clean_note_html, note_style)]], colWidths=[510])
            note_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#eef2ff")),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor("#a5b4fc")),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(note_table)
            story.append(Spacer(1, 10))

        # 4. Toàn Bộ Nội Dung Trích Xuất (Full Context)
        story.append(Paragraph("2. Toàn Bộ Nội Dung Chi Tiết (Full Extracted Content)", h1_style))
        
        # Tách từng đoạn văn bản
        for p_str in (ctx or "").split("\n\n"):
            clean = p_str.strip()
            if not clean:
                continue
            # Làm sạch thẻ ký tự đặc biệt cho ReportLab Paragraph
            safe_text = clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
            if safe_text.startswith("#"):
                clean_heading = safe_text.lstrip("#").strip()
                story.append(Paragraph(f"<b>{clean_heading}</b>", h1_style))
            else:
                story.append(Paragraph(safe_text, body_style))

        doc.build(story)
        filename = f"{entry_id}_{cls._sanitize_filename(title)}.pdf"
        return buf.getvalue(), "application/pdf", filename

    @classmethod
    def export_by_format(cls, entry: Dict[str, Any], fmt: str) -> Tuple[bytes, str, str]:
        """Hàm điều phối xuất theo định dạng: docx, xlsx, pdf, csv, json, txt."""
        fmt = (fmt or "docx").lower().strip()
        if fmt in ("word", "docx", "doc"):
            return cls.export_to_docx(entry)
        elif fmt in ("excel", "xlsx", "xls"):
            return cls.export_to_excel(entry)
        elif fmt in ("pdf",):
            return cls.export_to_pdf(entry)
        elif fmt in ("csv",):
            return cls.export_to_csv(entry)
        elif fmt in ("json",):
            return cls.export_to_json(entry)
        else:
            # Mặc định TXT
            entry_id = entry.get("id") or entry.get("ID", "RES-0000")
            title = entry.get("title") or entry.get("TITLE", "Untitled")
            ctx = entry.get("ctx") or entry.get("CTX", "")
            filename = f"{entry_id}_{cls._sanitize_filename(title)}.txt"
            return ctx.encode("utf-8"), "text/plain; charset=utf-8", filename
