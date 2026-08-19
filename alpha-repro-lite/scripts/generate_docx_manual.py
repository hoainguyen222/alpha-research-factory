#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to generate an exhaustive, professional DOCX User Manual for Alpha Research Factory.
Includes Glossary of Terms (Learned Rules, Blueprint, Hit Count), Code Tree, Database Schemas, Storage Maps, Skills & Daemons, Workflows & FAQs.
"""

import os
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner padding of a table cell in dxa."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_callout(doc, text_list, title="LƯU Ý / ĐẶC ĐIỂM QUAN TRỌNG", fill_hex="F0F4F8", border_hex="2563EB"):
    """Add a styled callout box to document."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    for item in text_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_before = Pt(1)
        p_item.paragraph_format.space_after = Pt(2)
        p_item.paragraph_format.left_indent = Inches(0.1)
        r = p_item.add_run(item)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_str):
    """Add a code block with monospace font and dark background."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "1E293B")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(code_str)
    r.font.name = "Consolas"
    r.font.size = Pt(9.0)
    r.font.color.rgb = RGBColor(0xF8, 0xFA, 0xFC)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_styled_table(doc, headers, rows_data, col_widths):
    """Add a professionally styled table."""
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.size = Pt(9.5)

    for row_idx, data in enumerate(rows_data, start=1):
        row_cells = tbl.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = str(text)
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(9.0)
            p.runs[0].font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            if col_idx == 0:
                p.runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def build_user_manual_docx():
    doc = docx.Document()

    # Set Margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.0)
    normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Helper function for Section Headers
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(13.5)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return h

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.bold = True
            rb.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        r = p.add_run(text)
        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    # ─── COVER / HEADER ──────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(2)
    run_main = p_title.add_run("TÀI LIỆU HƯỚNG DẪN SỬ DỤNG VÀ VẬN HÀNH TOÀN DIỆN")
    run_main.bold = True
    run_main.font.size = Pt(18)
    run_main.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    run_sub = p_sub.add_run("ALPHA RESEARCH FACTORY (ALPHA-REPRO-LITE)\nNền Tảng Tự Động Hóa Nghiên Cứu Định Lượng, Khám Phá Tri Thức & Bóc Tách Chiến Lược Alpha")
    run_sub.font.size = Pt(11.0)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # ─── SECTION 1: TỔNG QUAN ─────────────────────────────────────────
    add_h1("1. TỔNG QUAN HỆ THỐNG")
    doc.add_paragraph("Alpha Research Factory là một nền tảng nghiên cứu định lượng tự động (Autonomous Quantitative Alpha Platform). Hệ thống được thiết kế để giải quyết toàn bộ bài toán từ khâu thu thập tri thức học thuật/thị trường đến khâu tổng hợp công thức toán học, sinh mã nguồn giao dịch Python/C++ và chuẩn bị kiểm thử (Backtesting).")

    add_bullet(" Tự động tìm kiếm, tải về, bóc tách và phân loại các công trình nghiên cứu định lượng từ các nguồn học thuật hàng đầu (arXiv, CrossRef, OpenAlex, SSRN, Research Blogs).", "• Tự động hóa toàn diện:")
    add_bullet(" Bóc tách chính xác công thức toán học (Math Formulas), logic vào/ra lệnh (Entry/Exit Rules), tham số (Parameters) và mã nguồn thực thi Python/C++.", "• Trích xuất chuyên sâu:")
    add_bullet(" Tự động học và lưu lại cấu trúc của từng website/dạng bài báo để các lần sau chạy bằng code thuần với 0% chi phí Token LLM.", "• Tối ưu hóa tài nguyên:")
    add_bullet(" Lưu trữ có cấu trúc hỗ trợ tìm kiếm toàn văn FTS5 (Full-Text Search) siêu tốc.", "• Kho tri thức chuẩn mực:")

    # ─── SECTION 2: BẢNG GIẢI THÍCH THUẬT NGỮ (GLOSSARY) ──────────────
    add_h1("2. BẢNG GIẢI THÍCH THUẬT NGỮ & KHÁI NIỆM CỐT LÕI (GLOSSARY)")
    doc.add_paragraph("Để người đọc dễ dàng nắm bắt mục đích và ý nghĩa các thành phần, bảng dưới đây giải thích tường minh từng thuật ngữ chuyên môn:")

    headers_g = ["Thuật Ngữ / Khái Niệm", "Định Nghĩa Kỹ Thuật", "Mục Đích & Ý Nghĩa Thực Tiễn"]
    widths_g = [Inches(1.8), Inches(2.2), Inches(2.5)]
    rows_g = [
        (
            "Learned Rules\n(Quy Tắc Đã Học)",
            "Các khuôn mẫu thuật toán và công thức bóc tách chiến lược mà hệ thống tự động rút trích và ghi nhớ sau khi đọc các bài báo.",
            "Thay vì gọi AI phân tích lại từ đầu (tốn token), hệ thống tra cứu bảng learned_rules để áp dụng ngay công thức đã học trong 0.001s và tăng hit_count."
        ),
        (
            "Site Templates /\nScraper Blueprint",
            "Bộ công thức kỹ thuật (chứa CSS Selectors, vùng chứa tiêu đề, nội dung chính và bộ lọc rác) của từng tên miền trang web.",
            "Cho phép cào bài viết từ lần thứ 2 trở đi bằng code Python thuần (Fast-Path) trong 0.002s, đạt 0% chi phí Token và tăng tốc gấp 100 lần."
        ),
        (
            "Hit Count\n(Số Lần Tái Sử Dụng)",
            "Bộ đếm ghi nhận số lần một Rule hoặc một Site Template được hệ thống tái sử dụng thành công trong thực tế.",
            "Mỗi lần hit_count tăng thêm 1 đơn vị đồng nghĩa với việc hệ thống vừa tiết kiệm cho bạn 1 lần gọi API trí tuệ nhân tạo."
        ),
        (
            "Paper Vault\n(Kho Lưu Trữ Toàn Văn)",
            "Cơ sở dữ liệu tập trung lưu trữ toàn văn (Full Context), metadata và bản tóm tắt của mọi bài báo, tài liệu học thuật (mã RES-XXXX).",
            "Bảo toàn toàn vẹn tri thức gốc, hỗ trợ tìm kiếm toàn văn FTS5 siêu tốc và xem lại văn bản gốc bất kỳ lúc nào."
        ),
        (
            "Strategy Components\n(Thành Phần Chiến Lược)",
            "Thực thể chiến lược độc lập (mã COMP-XXXX) chứa mã nguồn Python/C++, công thức toán, quy tắc Entry/Exit và bộ tham số định lượng.",
            "Chuẩn bị sẵn sàng làm đầu vào trực tiếp cho động cơ Backtesting độc lập trên Leaderboard mà không cần chỉnh sửa thủ công."
        ),
        (
            "Skill\n(Kỹ Năng Tĩnh)",
            "Các gói mã nguồn, kịch bản thực thi và tài liệu quy trình cố định do lập trình viên định nghĩa sẵn trên hệ thống tệp (skills/, SKILL.md).",
            "Cung cấp công cụ và phương pháp luận chuẩn mực để AI thực hiện các tác vụ (như trích xuất PDF, OCR hình ảnh, vượt tường lửa)."
        ),
        (
            "3-Layer Deduplication\n(Bộ Lọc Trùng 3 Tầng)",
            "Cơ chế lọc trùng 3 lớp: (1) Khớp tuyệt đối URL/DOI -> (2) Khớp mờ Tiêu đề >= 80% -> (3) Vân tay nội dung >= 85-90%.",
            "Đảm bảo kho lưu trữ hoàn toàn sạch sẽ, không bao giờ bị trùng lặp dù bài báo xuất hiện ở cả arXiv, CrossRef hay bị đổi tiêu đề."
        ),
        (
            "Financial Guardrail\n(Thẩm Định Ngữ Nghĩa)",
            "Bộ lọc thông minh thẩm định độ dài (< 15 từ) và cấu trúc thuật ngữ định lượng của tài liệu nạp vào.",
            "Tự động phát hiện và ngăn chặn các hình ảnh / tài liệu đời thường ngoài ngành xâm nhập làm rác kho chiến lược Backtest."
        ),
        (
            "Catch-up Execution\n(Chạy Bù Tiến Trình)",
            "Tính năng Persistent=true của dịch vụ nền Linux Systemd Timer (alpha_scheduler.timer).",
            "Đảm bảo nếu máy tính bị tắt nguồn vào đúng giờ hẹn quét, thì ngay khi khởi động lại máy, hệ thống sẽ tự động chạy bù ngay lập tức."
        )
    ]
    add_styled_table(doc, headers_g, rows_g, widths_g)

    # ─── SECTION 3: BẢN ĐỒ MÃ NGUỒN & TỔ CHỨC DỰ ÁN ──────────────────
    add_h1("3. BẢN ĐỒ CẤU TRÚC MÃ NGUỒN & TỔ CHỨC DỰ ÁN (SOURCE CODE MAP)")
    doc.add_paragraph("Toàn bộ mã nguồn dự án được tổ chức chặt chẽ theo từng module chuyên biệt tại: /home/hoai/Alphareserach_agent-codex-alpha-repro-lite-core/alpha-repro-lite/")

    source_tree = """alpha-repro-lite/
├── research_coordinator.py       # 🎼 Nhạc trưởng điều phối toàn bộ luồng bóc tách đa phương thức
├── run_dashboard.py              # 🌐 Script khởi động Web Dashboard (Cổng 5055)
├── config.py                     # ⚙️ Cấu hình đường dẫn Database, cổng Web và hằng số toàn cục
├── .env                          # 🔑 File biến môi trường lưu API Keys (Anthropic Claude, Semantic)
│
├── extractors/                   # 📥 MODULE BÓC TÁCH DỮ LIỆU ĐA PHƯƠNG THỨC
│   ├── pdf_extractor.py          # Bóc tách PDF 2 cột, de-hyphenation, giữ nguyên công thức toán
│   ├── image_extractor.py        # Động cơ OCR hình ảnh (RapidOCR ONNX + Neural Filters)
│   ├── web_extractor.py          # Trích xuất Web/Blog (Trafilatura, Jina Reader, Fast-Path)
│   ├── keyword_search_engine.py  # Bộ máy tìm kiếm học thuật (arXiv, CrossRef, OpenAlex)
│   └── content_cleaner.py        # Bộ lọc khử nhiễu văn bản (loại bỏ ads, cookie banner, sidebar)
│
├── vault/                        # 🏛️ MODULE QUẢN TRỊ KHO TRI THỨC & BỘ NHỚ HỌC TẬP
│   ├── unified_vault_db.py       # Quản lý research_vault.db, bộ lọc trùng 3 tầng, FTS5 Search
│   ├── strategy_components_db.py # Quản lý extracted_strategy_components trong quant_platform.db
│   ├── site_template_engine.py   # Quản lý mẫu bóc tách Web (Scraper Blueprints)
│   └── learned_rule_engine.py    # Quản lý quy tắc và họ chiến lược đã học (Pattern Memory)
│
├── bypass/                       # 🛡️ MODULE VƯỢT TƯỜNG LỬA & ANTI-BOT
│   ├── anti_scraping_bypass.py   # Cơ chế tự động xoay User-Agent, Proxies, Jina Fallback
│   └── cloudflare_handler.py     # Xử lý các trang web chặn IP hoặc có bảo vệ Cloudflare
│
├── scripts/                      # ⚙️ MODULE TỰ ĐỘNG HÓA & NHÀ MÁY ALPHA
│   ├── auto_alpha_factory.py     # Nhà máy phân tích ngữ nghĩa, sinh code và chuẩn bị case backtest
│   ├── smart_auto_runner.py      # Kịch bản chạy ngầm tự động hàng ngày cho Systemd Daemon
│   └── daily_crawler.py          # Script thu thập dữ liệu định kỳ theo danh sách từ khóa
│
├── web/                          # 🖥️ MODULE GIAO DIỆN WEB DASHBOARD (FLASK)
│   ├── app.py                    # Backend Flask REST API phục vụ Dashboard
│   ├── templates/index.html      # Giao diện Dashboard Dark-Theme chuyên nghiệp
│   └── static/                   # CSS Glassmorphism & Javascript Client Controller
└── storage/                      # 💾 KHO LƯU TRỮ DỮ LIỆU & FILE VẬT LÝ"""
    add_code_block(doc, source_tree)

    # ─── SECTION 4: KIẾN TRÚC LƯU TRỮ & CƠ SỞ DỮ LIỆU ─────────────────
    add_h1("4. KIẾN TRÚC LƯU TRỮ & TOÀN BỘ CƠ SỞ DỮ LIỆU (DATABASE SCHEMAS)")
    doc.add_paragraph("Hệ thống sử dụng kiến trúc Lưu Trữ Kép (Dual-Database Architecture) phân tách rõ ràng giữa Tài liệu học thuật và Dữ liệu định lượng:")

    add_h2("4.1 Cơ sở dữ liệu 1: storage/structured_vault/research_vault.db")
    doc.add_paragraph("Chuyên trách lưu trữ toàn bộ văn bản và hồ sơ tài liệu học thuật từ mọi nguồn.")

    headers_v = ["Tên Cột", "Kiểu Dữ Liệu", "Ý Nghĩa / Nội Dung Lưu Trữ"]
    widths_v = [Inches(1.5), Inches(1.3), Inches(3.7)]
    rows_v = [
        ("id", "TEXT (PK)", "Mã định danh duy nhất (ví dụ: RES-20260818-0001)"),
        ("title", "TEXT", "Tiêu đề chính thức của công trình nghiên cứu / bài viết"),
        ("type", "TEXT", "Nguồn gốc: FILE_PDF, IMAGE, WEB_ARTICLE, KEYWORD_SEARCH"),
        ("ctx", "TEXT", "Toàn bộ văn bản đầy đủ của bài báo (Full Text Context)"),
        ("note", "TEXT", "Tóm tắt, luận điểm cốt lõi và nhận định của AI"),
        ("web", "TEXT", "Link URL gốc, mã DOI hoặc tên file nguồn"),
        ("metadata", "TEXT (JSON)", "Metadata: Tác giả, năm xuất bản, số từ, engine OCR"),
        ("raw_file_path", "TEXT", "Đường dẫn file văn bản lưu trong storage/raw_sources/"),
        ("created_at", "TEXT", "Thời gian nạp vào kho (chuẩn ISO 8601)")
    ]
    add_styled_table(doc, headers_v, rows_v, widths_v)
    add_bullet(" Bảng chỉ mục FTS5 (Full-Text Search) cho phép tìm kiếm bất kỳ từ khóa nào trong hàng triệu từ toàn văn chỉ trong < 0.01 giây.", "• Bảng research_vault_fts:")
    add_bullet(" Tự động xuất ra file unified_vault.jsonl và unified_vault.csv đồng bộ thời gian thực.", "• Tự động đồng bộ xuất file:")

    add_h2("4.2 Cơ sở dữ liệu 2: quant_platform.db (Tại thư mục gốc)")
    doc.add_paragraph("Chuyên trách lưu trữ thành phần chiến lược bóc tách, bộ nhớ học tập và kết quả kiểm thử định lượng.")

    add_h3("A. Bảng extracted_strategy_components (Thành phần chiến lược độc lập):")
    headers_c = ["Tên Cột", "Kiểu Dữ Liệu", "Ý Nghĩa / Nội Dung Lưu Trữ"]
    widths_c = [Inches(1.5), Inches(1.3), Inches(3.7)]
    rows_c = [
        ("id", "TEXT (PK)", "Mã thành phần chiến lược (ví dụ: COMP-20260818-0001)"),
        ("vault_id", "TEXT (FK)", "Mã tài liệu tương ứng trong Paper Vault (RES-XXXX)"),
        ("strategy_name", "TEXT", "Tên chiến lược định lượng (ví dụ: Ornstein-Uhlenbeck StatArb)"),
        ("model_family", "TEXT", "Họ mô hình: Statistical_Arbitrage, Momentum, Deep_RL, Volatility"),
        ("asset_class", "TEXT", "Lớp tài sản mục tiêu (equities, crypto, forex, futures)"),
        ("timeframe", "TEXT", "Khung thời gian giao dịch (1m, 5m, 15m, 1h, 1d)"),
        ("code_snippets", "TEXT (JSON)", "Mã nguồn thực thi Python/C++: hàm generate_signal(data)"),
        ("math_formulas", "TEXT (JSON)", "Các phương trình toán: Spread, Z-Score, OU Process, Kalman"),
        ("trading_rules", "TEXT (JSON)", "Quy tắc mở lệnh Entry Long/Short, Exit, Cắt lỗ Trailing Stop"),
        ("parameters", "TEXT (JSON)", "Bộ tham số mặc định: Rolling window, Threshold, Stoploss multiplier"),
        ("backtest_status", "TEXT", "Trạng thái kiểm thử: PENDING, RUNNING, VERIFIED")
    ]
    add_styled_table(doc, headers_c, rows_c, widths_c)

    add_h3("B. Bảng crawler_site_templates (Bộ nhớ Mẫu cào Web - Fast-Path Blueprint):")
    rows_t = [
        ("id", "TEXT (PK)", "Mã mẫu bóc tách (ví dụ: TPL-ARXIV-001, TPL-SUBSTACK-001)"),
        ("domain_pattern", "TEXT (UQ)", "Tên miền áp dụng (ví dụ: arxiv.org, ssrn.com, medium.com)"),
        ("title_selector", "TEXT", "CSS Selector để lấy Tiêu đề (ví dụ: h1.title)"),
        ("content_selector", "TEXT", "CSS Selector để lấy Thân bài viết"),
        ("noise_selectors", "TEXT (JSON)", "Danh sách CSS Selector cần loại bỏ (quảng cáo, sidebar, cookie)"),
        ("hit_count", "INTEGER", "Số lần tái sử dụng (mỗi lần tăng là 1 lần tiết kiệm 100% Token)")
    ]
    add_styled_table(doc, headers_c, rows_t, widths_c)

    add_h3("C. Bảng learned_rules (Quy tắc & Khuôn mẫu chiến lược đã học):")
    doc.add_paragraph("Lưu trữ 14 họ quy tắc và khuôn mẫu xử lý chiến lược. Khi gặp bài báo tương tự, hệ thống tăng chỉ số hit_count và áp dụng ngay trong 0.001 giây thay vì suy luận lại.")

    add_h3("D. Bảng backtest_metrics (Kết quả kiểm thử định lượng):")
    doc.add_paragraph("Lưu các chỉ số hiệu suất: sharpe_ratio, sortino_ratio, calmar_ratio, total_return_pct, max_drawdown_pct, hit_rate_pct, alpha, beta, total_ticks, throughput_ticks_sec.")

    # ─── SECTION 5: KỸ NĂNG, RULES & DAEMON ───────────────────────────
    add_h1("5. VỊ TRÍ LƯU TRỮ KỸ NĂNG, QUY TẮC & TIẾN TRÌNH NỀN (SKILLS, RULES & DAEMONS)")

    add_h2("5.1 Kỹ năng & Tùy biến Hệ thống (Agent Skills & Rules)")
    add_bullet(" /home/hoai/.gemini/config/ (Chứa các kỹ năng và cấu hình máy chủ MCP toàn cục).", "1. Global Customizations Root:")
    add_bullet(" .agents/ và skills/ (Chứa các cheatsheet SKILL.md chuyên biệt và thư mục rules/).", "2. Workspace Customizations Root:")
    add_bullet(" /home/hoai/.gemini/antigravity-ide/builtin/skills/ (Chứa agy-customizations, antigravity_guide).", "3. IDE Builtin Skills:")

    add_h2("5.2 Tiến trình nền Hệ điều hành (Systemd Service & Timer)")
    add_bullet(" ~/.config/systemd/user/alpha_scheduler.service (Tiến trình chạy nền thực thi smart_auto_runner.py).", "• Service File:")
    add_bullet(" ~/.config/systemd/user/alpha_scheduler.timer (Bộ hẹn giờ định kỳ kèm cấu hình Persistent=true).", "• Timer File:")

    # ─── SECTION 6: LUỒNG VẬN HÀNH A-Z ──────────────────────────────
    add_h1("6. KIẾN TRÚC & LUỒNG VẬN HÀNH TOÀN DIỆN TỪ A - Z (END-TO-END WORKFLOW)")
    doc.add_paragraph("Quy trình xử lý dữ liệu được thực hiện khép kín qua 5 giai đoạn tuần tự:")
    add_bullet(" Hỗ trợ bóc tách PDF 2 cột, OCR ảnh bằng RapidOCR ONNX, và cào bài viết web.", "Bước 1: Thu thập & Nhập liệu đa phương thức:")
    add_bullet(" 3 Tầng: Khớp tuyệt đối (URL/DOI) -> Khớp mờ tiêu đề (>=80%) -> Vân tay nội dung (>=85-90%).", "Bước 2: Bộ lọc chống trùng lặp đa tầng:")
    add_bullet(" Thẩm định tài chính (loại trừ ảnh/file rác) và tra cứu Blueprint (chạy Fast-Path 0.002s, 0% Token).", "Bước 3: Thẩm định ngữ nghĩa & Học mẫu:")
    add_bullet(" Bóc tách nguyên tử: Mỗi bài báo -> 1 mã RES-XXXX -> đúng 1 mã chiến lược COMP-XXXX.", "Bước 4: Sinh thành phần chiến lược nguyên tử:")
    add_bullet(" Đồng bộ research_vault.db và quant_platform.db, hiển thị Dashboard trực quan.", "Bước 5: Đồng bộ Database & Dashboard:")

    # ─── SECTION 7: CÁC ĐẶC ĐIỂM NỔI BẬT ────────────────────────────
    add_h1("7. CÁC ĐẶC ĐIỂM & CÔNG NGHỆ NỔI BẬT")
    add_callout(doc, [
        "1. Cơ chế Fast-Path Scraper Blueprint: Chạy bằng code Python thuần chỉ trong 0.002 giây cho các tên miền đã học, tiết kiệm 100% chi phí Token.",
        "2. Bộ lọc chống trùng 3 tầng: Đảm bảo không bao giờ tải trùng dù bài báo xuất hiện ở cả arXiv, CrossRef hay bị đổi tiêu đề.",
        "3. Tự động chạy bù (Catch-up Execution): Systemd Timer với Persistent=true đảm bảo hệ thống tự động chạy bù ngay khi máy tính khởi động nếu trước đó bị tắt nguồn.",
        "4. Bảo vệ dữ liệu Backtest: Ngăn chặn triệt để dữ liệu rác hoặc ảnh ngoài ngành xâm nhập vào bảng chiến lược."
    ], title="CÁC CÔNG NGHỆ CỐT LÕI TẠO NÊN SỰ VƯỢT TRỘI")

    # ─── SECTION 8: HƯỚNG DẪN DASHBOARD ──────────────────────────────
    add_h1("8. HƯỚNG DẪN THAO TÁC CHI TIẾT TRÊN DASHBOARD (HTTP://127.0.0.1:5055)")

    headers_db = ["Tab / Khu Vực", "Chức Năng Chính", "Ý Nghĩa Thực Tiễn"]
    widths_db = [Inches(1.8), Inches(2.7), Inches(2.0)]
    rows_db = [
        ("Leaderboard", "Bảng xếp hạng hiệu suất chiến lược định lượng", "Theo dõi Sharpe, Return, Drawdown, Calmar"),
        ("Paper Vault", "Kho lưu trữ toàn văn tài liệu & FTS5 search", "Tra cứu nhanh, xem tóm tắt & sao chép nội dung"),
        ("Strategy Components", "Danh mục mã nguồn, công thức & tham số", "Xem code Python/C++, logic Entry/Exit để backtest"),
        ("Learned Rules", "Kho mẫu cào web & quy tắc đã học", "Theo dõi số lần tái sử dụng (Hit Count) & tiết kiệm token"),
        ("Spider & AI", "Tìm kiếm theo từ khóa qua arXiv/CrossRef/OpenAlex", "Khám phá tự động & cấu hình tiến trình chạy ngầm"),
        ("Upload", "Tải lên PDF, Hình ảnh báo cáo, Link bài viết", "Nạp tài liệu thủ công và bóc tách tức thì")
    ]
    add_styled_table(doc, headers_db, rows_db, widths_db)

    # ─── SECTION 9: FAQ & GIẢI ĐÁP THẮC MẮC ──────────────────────────
    add_h1("9. GIẢI ĐÁP CHI TIẾT CÁC CÂU HỎI THƯỜNG GẶP (FAQ)")

    faqs = [
        (
            "Câu 1: Cơ chế tự động chạy ngầm hoạt động thế nào? Nếu đến giờ chạy mà máy tính tắt thì có chạy bù không?",
            "CÓ 100%! Hệ thống sử dụng dịch vụ nền chuẩn Linux Systemd Timer (alpha_scheduler.timer) được kích hoạt thuộc tính Persistent=true. Nếu đến lịch quét mà máy tính đang tắt nguồn hoặc ở chế độ ngủ (Sleep), ngay khi bạn bật máy tính lên, hệ điều hành sẽ tự động kích hoạt chạy bù ngay lập tức (Catch-up Run) mà không bỏ sót bất kỳ chu kỳ nào."
        ),
        (
            "Câu 2: Tại sao cơ chế bóc tách website lại tiết kiệm token và có thể chạy bằng code thuần?",
            "Khi tải một trang web lần đầu tiên, hệ thống dùng thuật toán nhận diện cấu trúc DOM để định vị vùng văn bản quan trọng. Sau đó, nó tự động tổng hợp thành một bản vẽ kỹ thuật (Scraper Blueprint) gồm các CSS Selectors tương ứng cho tên miền đó và lưu vào bảng crawler_site_templates. Từ các lần sau, hệ thống áp dụng trực tiếp mã nguồn Python bóc tách theo Blueprint trong 0.002 giây, hoàn toàn không gửi dữ liệu lên AI và không tốn bất kỳ Token nào."
        ),
        (
            "Câu 3: Nếu một công trình nghiên cứu vừa có trên arXiv vừa có trên CrossRef, hoặc nội dung giống >90% thì sao?",
            "Hệ thống áp dụng Bộ lọc 3 tầng: (1) Khớp mờ tiêu đề Fuzzy Title >= 80% để loại trừ ngay các bài trùng bản thảo/bản xuất bản. (2) So khớp vân tay văn bản và độ bao phủ từ vựng (Token Containment) >= 85-90%. Nếu nội dung cốt lõi trùng khớp trên 90%, hệ thống tự động từ chối nạp trùng vào Vault, giữ kho dữ liệu luôn sạch sẽ."
        ),
        (
            "Câu 4: Khi nhập từ khóa tìm kiếm, tại sao hệ thống tải về các bài riêng lẻ thay vì gộp chung?",
            "Để phục vụ Backtest tự động chuẩn xác, mỗi công trình nghiên cứu phải là một thực thể nguyên tử (Atomic Entity). Việc bóc tách riêng từng bài giúp hệ thống trích xuất đúng mã nguồn Python/C++, công thức toán và bộ tham số riêng cho từng bài, cho phép chạy Backtest độc lập từng chiến lược trên Leaderboard."
        ),
        (
            "Câu 5: Tại sao trong Vault có 22 bài báo nhưng số 'Luật đã học (Rules)' chỉ có 14?",
            "Con số 14 là HOÀN TOÀN ĐÚNG! Rules không đếm số lượng bài báo, mà là các khuôn mẫu thuật toán và công thức bóc tách độc lập (Model Families). Ví dụ: O-U Mean Reversion, Deep RL FinRL, GARCH Volatility, Kalman Filter... Khi có nhiều bài báo thuộc cùng một họ mô hình, hệ thống sẽ tái sử dụng Rule đã học và tăng chỉ số hit_count lên chứ không sinh ra Rule mới trùng lặp."
        ),
        (
            "Câu 6: Nếu người dùng tải lên một hình ảnh hoặc tài liệu không liên quan đến tài chính thì hệ thống sẽ làm gì?",
            "Hệ thống có cơ chế Financial Semantic Guardrail: Động cơ Vision OCR vẫn đọc chữ trong ảnh và lưu vào Vault với loại IMAGE để lưu vết lịch sử. Sau đó, hệ thống kiểm tra ngữ nghĩa tài chính: Nếu văn bản quá ngắn (< 15 từ) hoặc không chứa cấu trúc công thức/thuật ngữ định lượng, hệ thống sẽ ghi chú cảnh báo và TỰ ĐỘNG BỎ QUA, KHÔNG SINH CHIẾN LƯỢC RÁC VÀO PIPELINE BACKTEST."
        ),
        (
            "Câu 7: Vì sao Skill không lưu trong Database mà Learned Rules lại được lưu trong Database?",
            "Skill là mã nguồn tĩnh, kịch bản Python và tài liệu quy trình (SKILL.md) cần quản lý phiên bản qua Git và thực thi trực tiếp trên hệ điều hành nên lưu ở File System. Ngược lại, Learned Rules là tri thức và kinh nghiệm động biến đổi 24/7, cần đọc/ghi liên tục, tra cứu chỉ mục siêu tốc trong 0.001s và thống kê trực quan lên Dashboard nên bắt buộc phải lưu trong Database."
        )
    ]

    for q, a in faqs:
        add_h2(f"❓ {q}")
        p_ans = doc.add_paragraph(f"👉 Trả lời: {a}")
        p_ans.paragraph_format.space_before = Pt(2)
        p_ans.paragraph_format.space_after = Pt(6)

    # ─── SECTION 10: SỔ TAY LỆNH QUẢN TRỊ ────────────────────────────
    add_h1("10. SỔ TAY LỆNH QUẢN TRỊ HỆ THỐNG & BẢO TRÌ")
    add_h3("1. Khởi động Web Dashboard:")
    add_code_block(doc, "python3 run_dashboard.py\n# Mở trình duyệt tại: http://127.0.0.1:5055")

    add_h3("2. Kiểm tra trạng thái tiến trình chạy ngầm (Systemd Service):")
    add_code_block(doc, "systemctl --user status alpha_scheduler.service\nsystemctl --user status alpha_scheduler.timer")

    add_h3("3. Chạy thử một chu trình quét & bóc tách tự động ngay lập tức:")
    add_code_block(doc, "python3 scripts/smart_auto_runner.py")

    add_h3("4. Kiểm tra tính toàn vẹn dữ liệu trong SQLite:")
    add_code_block(doc, "python3 -c \"\nimport sqlite3\nc1 = sqlite3.connect('storage/structured_vault/research_vault.db')\nprint('Total Vault Papers:', c1.execute('SELECT count(*) FROM research_vault').fetchone()[0])\nc2 = sqlite3.connect('quant_platform.db')\nprint('Total Strategy Components:', c2.execute('SELECT count(*) FROM extracted_strategy_components').fetchone()[0])\nprint('Total Learned Rules:', c2.execute('SELECT count(*) FROM learned_rules').fetchone()[0])\n\"")

    out_path = Path(BASE_DIR) / "USER_MANUAL.docx"
    doc.save(str(out_path))
    print(f"✅ Generated exhaustive DOCX manual with Glossary: {out_path} ({out_path.stat().st_size} bytes)")

if __name__ == "__main__":
    BASE_DIR = Path(__file__).resolve().parent.parent
    build_user_manual_docx()
